from django.http import HttpResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import io

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_bookings_excel(request):
    """Экспорт данных профиля в Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Профиль"
    
    # Заголовки
    headers = ['Поле', 'Значение']
    ws.append(headers)
    
    # Стилизация заголовков
    header_fill = PatternFill(start_color="E7DC00", end_color="E7DC00", fill_type="solid")
    header_font = Font(bold=True, size=12)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
    
    # Данные пользователя
    ws.append(['Имя пользователя', request.user.username])
    ws.append(['Email', request.user.email])
    ws.append(['Имя', request.user.first_name or 'Не указано'])
    ws.append(['Фамилия', request.user.last_name or 'Не указано'])
    
    try:
        profile = request.user.userprofile
        ws.append(['Телефон', profile.phone or 'Не указан'])
        ws.append(['Дата рождения', str(profile.birth_date) if profile.birth_date else 'Не указана'])
    except:
        pass
    
    # Автоширина колонок
    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column[0].column_letter].width = adjusted_width
    
    # Сохранение
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=profile.xlsx'
    return response

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_bookings_word(request):
    """Экспорт профиля в Word"""
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Мой профиль', 0)
    title.alignment = 1
    
    # Информация о пользователе
    doc.add_paragraph(f'Имя пользователя: {request.user.username}')
    doc.add_paragraph(f'Email: {request.user.email}')
    doc.add_paragraph(f'Имя: {request.user.first_name or "Не указано"}')
    doc.add_paragraph(f'Фамилия: {request.user.last_name or "Не указано"}')
    
    try:
        profile = request.user.userprofile
        doc.add_paragraph(f'Телефон: {profile.phone or "Не указан"}')
        doc.add_paragraph(f'Дата рождения: {profile.birth_date or "Не указана"}')
    except:
        pass
    
    # Сохранение
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=profile.docx'
    return response

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_bookings_pdf(request):
    """Экспорт профиля в PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#E7DC00'),
        spaceAfter=30,
        alignment=1
    )
    
    # Заголовок
    elements.append(Paragraph('Мой профиль', title_style))
    elements.append(Spacer(1, 12))
    
    # Информация о пользователе
    user_info = f"""Имя пользователя: {request.user.username}<br/>
    Email: {request.user.email}<br/>
    Имя: {request.user.first_name or 'Не указано'}<br/>
    Фамилия: {request.user.last_name or 'Не указано'}"""
    
    try:
        profile = request.user.userprofile
        user_info += f"<br/>Телефон: {profile.phone or 'Не указан'}<br/>"
        user_info += f"Дата рождения: {profile.birth_date or 'Не указана'}"
    except:
        pass
    
    elements.append(Paragraph(user_info, styles['Normal']))
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=profile.pdf'
    return response

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_menu_excel(request):
    """Экспорт меню в Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Меню"
    
    headers = ['Название', 'Описание', 'Цена']
    ws.append(headers)
    
    header_fill = PatternFill(start_color="E7DC00", end_color="E7DC00", fill_type="solid")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True)
    
    # Примеры блюд
    menu_data = [
        ['СЕТ ИЗ КОЛБАСОК К ПИВУ', 'Куриные колбаски, свино-говяжьи, охотничьи колбаски с соусами на выбор (410гр.)', 1100],
        ['ПЕРВЫЙ ШАГ Б/А', 'Безалкогольный лагер янтарно-золотистого цвета с лёгкой хмелевой горчинкой (0.5л)', 350],
    ]
    
    for item in menu_data:
        ws.append(item)
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=menu.xlsx'
    return response